from typing import List
from dataclasses import dataclass, field
from xml.dom import NotFoundErr
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

@dataclass
class FormatFilter:
    bold: List or str = field(default_factory = lambda: ["None", "True"])
    italic: List or str = field(default_factory = lambda: ["None", "True"])
    text: str or None = None
    font_rgb_color: List or str = "None"
    style_name: List or str = field(default_factory = lambda: ["Normal", "No Spacing"])

class ReadDocFile:
    def __init__(self, filename):
        self.fname = filename
        self.loaded_document = None
        self.content = None

    def load_document(self):
        try:
            self.loaded_document = Document(self.fname)
            self.content = self.get_content_as_list()
        except PackageNotFoundError:
            print(f"File not found: {self.fname}")
            quit()

    def get_content_as_list(self):
        return [paragraph.text for paragraph in self.loaded_document.paragraphs]

    def _get_bold(self, paragraph, filter):
        filter = [filter] if isinstance(filter, str) else filter

        for run in paragraph.runs:
            if str(run.bold) in filter:
                return paragraph

    def _get_italic(self, paragraph, filter):
        filter = [filter] if isinstance(filter, str) else filter

        for run in paragraph.runs:
            if str(run.italic) in filter:
                return paragraph

    def _get_exact_text(self, paragraph, filter):
        if paragraph.text == filter or filter is None:
            return paragraph

    def _get_style_name(self, paragraph, filter):
        filter = [filter] if isinstance(filter, str) else filter

        if str(paragraph.style.name) in filter:
            return paragraph

    def _get_font_color(self, paragraph, filter):
        filter = [filter] if isinstance(filter, str) else filter

        for run in paragraph.runs:
            if str(run.font.color.rgb) in str(filter):
                return paragraph

    def get_formatted_text(self, format_filter: FormatFilter):
        match_list = []
        for paragraph in self.loaded_document.paragraphs:
            bold = self._get_bold(paragraph, format_filter.bold)
            italic = self._get_italic(paragraph, format_filter.italic)
            exact_text = self._get_exact_text(paragraph, format_filter.text)
            style_name = self._get_style_name(paragraph, format_filter.style_name)
            font_color_rgb = self._get_font_color(paragraph, format_filter.font_rgb_color)

            if all([bold, italic, exact_text, style_name, font_color_rgb]):
                match_list.append(paragraph.text)
        return match_list

    def get_segment(self, exact_start="TBD", exact_end="TBD", contains_start="TBD", contains_end="TBD"):
        segment_list = []
        started = False
        for paragraph in self.loaded_document.paragraphs:
            if paragraph.text == exact_start or contains_start in paragraph.text:
                segment_list.append(paragraph.text)
                started = True
            elif paragraph.text == exact_end or contains_end in paragraph.text:
                segment_list.append(paragraph.text)
                return segment_list
            elif started:
                segment_list.append(paragraph.text)
        raise NotFoundErr(f"End text not found: {exact_end}")

    def get_table_content(self, table_index):
        table = self.loaded_document.tables[table_index]
        data = []
        keys = None

        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            if i == 0:
                keys = tuple(text)
                continue

            row_data = dict(zip(keys, text))
            data.append(row_data)

        return data


def docx_to_text():
    import docx2txt
    # extract text
    text1 = docx2txt.process("cv.docx")

    # extract text and write images in /tmp/img_dir
    # text2 = docx2txt.process("cv.docx", "YOUR_PATH")
    return text1


if __name__ == "__main__":
    doc = ReadDocFile("cv.docx")
    doc.load_document()

    # Find paragraph by filter
    filter = FormatFilter(bold="True", font_rgb_color="FF0000", italic="None")
    print(doc.get_formatted_text(format_filter=filter))

    # Find segment exact text
    result = doc.get_segment(exact_start="Objetivo", exact_end="")
    # print(result)

    # Find segment contains text
    result = doc.get_segment(contains_start="candidato", contains_end="celular")
    # print(result)

    # Get table content
    result = doc.get_table_content(table_index=0)
    print(result)

    # Use docx2txt
    print(docx_to_text())